#!/usr/bin/env python3
"""解析 iCloud 中的中英文 Word 简历，提取工作经历数据。

读取 context_CH.docx 和 context_EN.docx，解析每个经历的
period / title / company / bullets，输出为 JSON。
"""

from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path
from typing import Any

import docx

# ---- 配置 ----
ICLOUD_CV = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/private/CV"
PROJECT_ROOT = Path(__file__).resolve().parent.parent

# 现有 resume.ts 中保留的元数据（姓名、联系方式、教育、关于我）
# 这些字段不在 Word 文档中，同步时保持不变
META_ZH: dict[str, Any] = {
    "name": "蒋卢",
    "title": "市场总监",
    "phone": "18621971310",
    "email": "nickycom@msn.com",
    "education": {
        "period": "2003 – 2007",
        "degree": "管理学学士",
        "major": "信息管理与信息系统",
        "school": "上海大学",
    },
    "about": [
        "MBTI 人格类型: ENTJ",
        "国际精酿啤酒一级认证 (Cicerone)",
        "国际葡萄酒&烈酒二级认证 (WSET Level 2)",
    ],
}

META_EN: dict[str, Any] = {
    "name": "Nick Jiang",
    "title": "Marketing Director",
    "phone": "18621971310",
    "email": "nickycom@msn.com",
    "education": {
        "period": "2003 – 2007",
        "degree": "Bachelor of Business",
        "major": "Information Management & Systems",
        "school": "Shanghai University",
    },
    "about": [
        "MBTI: ENTJ",
        "Cicerone Global Beer Level 1 Certification",
        "WSET Level 2 Certification (Wine & Spirits)",
    ],
}


# ---- 解析逻辑 ----
def parse_experiences(doc: docx.Document) -> list[dict[str, Any]]:
    """从 Word 文档中提取所有工作经历。

    识别规则：
    - 经历以 "日期 标题 公司" 格式的段落开头
    - 后续非空段落为 bullet points
    - 空段落表示经历结束
    """
    experiences: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    # 日期模式：支持多种格式
    header_pattern = re.compile(
        r"^(\d{2}/\d{4})\s*[—\-–]\s*"  # start date
        r"(\d{2}/\d{4}|至今|present)",  # end date
        re.IGNORECASE,
    )

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # 空段落 = 当前经历结束
            if current and current.get("bullets"):
                experiences.append(current)
            current = None
            continue

        # 检查是否是新经历的标题行
        if header_pattern.match(text):
            # 保存上一个经历
            if current and current.get("bullets"):
                experiences.append(current)

            # 解析标题行：日期 + 职位 + 公司
            current = _parse_header(text)
        elif current is not None:
            # 当前经历的 bullet point
            text = re.sub(r"^[-•·]\s*", "", text)  # 去除前导符号
            current.setdefault("bullets", []).append(text)

    # 保存最后一个经历
    if current and current.get("bullets"):
        experiences.append(current)

    return experiences


def _parse_header(line: str) -> dict[str, Any]:
    """解析经历标题行，提取 period, title, company。

    格式示例：
    - "10/2024-至今            外部品牌顾问 外资啤酒品牌"
    - "10/2024 –present  Brand & Marketing Consultant , An Imported Beer Brand"
    - "08/2023 –09/2024     市场总监，华润雪花华南深圳营销中心"
    """
    # 统一空白字符：TAB → 空格
    line = line.replace("\t", " ")

    # 先提取日期部分
    date_match = re.match(
        r"(\d{2}/\d{4}\s*[—\-–]\s*(?:\d{2}/\d{4}|至今|present))\s+(.*)",
        line,
        re.IGNORECASE,
    )
    if not date_match:
        return {"period": line, "title": "", "company": "", "bullets": []}

    period = date_match.group(1).strip()
    rest = date_match.group(2).strip()

    # 多级分割策略：逗号 > 连续空格 > 单空格（中文兜底）
    if "，" in rest:
        parts = rest.split("，", 1)
    elif ", " in rest:
        parts = rest.split(", ", 1)
    elif "," in rest:
        parts = rest.split(",", 1)
    elif re.search(r"\s{2,}", rest):
        parts = re.split(r"\s{2,}", rest, maxsplit=1)
    elif " " in rest and not _has_ascii_upper(rest):
        # 纯中文无逗号场景：用第一个空格分割（如"外部品牌顾问 外资啤酒品牌"）
        parts = rest.split(" ", 1)
    else:
        parts = [rest]

    title = parts[0].strip()
    company = parts[1].strip() if len(parts) > 1 else ""

    return {"period": period, "title": title, "company": company, "bullets": []}


def _has_ascii_upper(text: str) -> bool:
    """检查文本是否包含 ASCII 大写字母（判断是否为英文内容）。"""
    return bool(re.search(r"[A-Z]{2,}", text))


# ---- 生成 TypeScript ----
def generate_resume_ts(
    experiences_zh: list[dict[str, Any]],
    experiences_en: list[dict[str, Any]],
) -> str:
    """生成 resume.ts 文件内容。"""

    def format_exps(exps: list[dict[str, Any]], indent: int = 8) -> str:
        items = []
        for exp in exps:
            bullets = ",\n".join(
                f'{" " * (indent + 4)}"{_escape(b)}"' for b in exp["bullets"]
            )
            item = (
                f'{" " * indent}{{\n'
                f'{" " * (indent + 2)}period: "{_escape(exp["period"])}",\n'
                f'{" " * (indent + 2)}title: "{_escape(exp["title"])}",\n'
                f'{" " * (indent + 2)}company: "{_escape(exp["company"])}",\n'
                f'{" " * (indent + 2)}bullets: [\n{bullets}\n{" " * (indent + 2)}],\n'
                f'{" " * indent}}}'
            )
            items.append(item)
        return ",\n".join(items)

    def format_meta(meta: dict[str, Any], indent: int = 6) -> str:
        about = ",\n".join(
            f'{" " * (indent + 2)}"{_escape(a)}"' for a in meta["about"]
        )
        edu = meta["education"]
        return (
            f'{" " * indent}name: "{_escape(meta["name"])}",\n'
            f'{" " * indent}title: "{_escape(meta["title"])}",\n'
            f'{" " * indent}phone: "{_escape(meta["phone"])}",\n'
            f'{" " * indent}email: "{_escape(meta["email"])}",\n'
            f'{" " * indent}education: {{\n'
            f'{" " * (indent + 2)}period: "{_escape(edu["period"])}",\n'
            f'{" " * (indent + 2)}degree: "{_escape(edu["degree"])}",\n'
            f'{" " * (indent + 2)}major: "{_escape(edu["major"])}",\n'
            f'{" " * (indent + 2)}school: "{_escape(edu["school"])}",\n'
            f'{" " * indent}}},\n'
            f'{" " * indent}about: [\n{about}\n{" " * indent}],'
        )

    return f"""export type Experience = {{
  period: string;
  title: string;
  company: string;
  bullets: string[];
}};

export type ResumeData = {{
  name: string;
  title: string;
  phone: string;
  email: string;
  education: {{
    period: string;
    degree: string;
    major: string;
    school: string;
  }};
  about: string[];
  experiences: Experience[];
}};

const data: Record<"zh" | "en", ResumeData> = {{
  zh: {{
{format_meta(META_ZH)}
    experiences: [
{format_exps(experiences_zh)}
    ],
  }},

  en: {{
{format_meta(META_EN)}
    experiences: [
{format_exps(experiences_en)}
    ],
  }},
}};

export default data;
"""


def _escape(s: str) -> str:
    """转义字符串中的特殊字符。"""
    return s.replace("\\", "\\\\").replace('"', '\\"')


# ---- 主流程 ----
def main() -> Path:
    zh_path = ICLOUD_CV / "context_CH.docx"
    en_path = ICLOUD_CV / "context_EN.docx"

    if not zh_path.exists():
        print(f"❌ 中文简历不存在: {zh_path}")
        sys.exit(1)
    if not en_path.exists():
        print(f"❌ 英文简历不存在: {en_path}")
        sys.exit(1)

    print(f"📄 解析中文简历: {zh_path.name}")
    doc_zh = docx.Document(str(zh_path))
    exps_zh = parse_experiences(doc_zh)
    print(f"   → 提取 {len(exps_zh)} 段工作经历")

    print(f"📄 解析英文简历: {en_path.name}")
    doc_en = docx.Document(str(en_path))
    exps_en = parse_experiences(doc_en)
    print(f"   → 提取 {len(exps_en)} 段工作经历")

    # 生成 resume.ts
    ts_content = generate_resume_ts(exps_zh, exps_en)
    output_path = PROJECT_ROOT / "src" / "data" / "resume.ts"
    output_path.write_text(ts_content, encoding="utf-8")
    print(f"✅ 已写入: {output_path}")

    return output_path


if __name__ == "__main__":
    main()
